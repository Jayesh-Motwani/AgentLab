from api.ai.schemas import ResearchResult, ReportFile  # re-import to get types
from pathlib import Path
import requests
from langchain_core.tools import tool
from api.myemailer.sender import send_mail
from api.myemailer.inbox_reader import read_inbox
from api.ai.services import generate_email_message
from api.ai.llms import get_openai_llm
from api.ai.schemas import Source, SearchQueries
from parallel import Parallel
import os
from urllib.parse import urlparse
import arxiv
from docx import Document
from api.ai.research_context import current_sources


PARALLEL_API_KEY = os.environ.get("PARALLEL_API_KEY")


def _publisher_from_url(url: str) -> str:
    """Derive a human-readable publisher name from the domain."""
    netloc = urlparse(url).netloc
    netloc = netloc.replace("www.", "")
    # e.g. "reuters.com" -> "Reuters"
    domain = netloc.split(".")[0]
    return domain.capitalize()


@tool
def send_me_email(subject: str, content: str, attachment_path: str | None = None, to_email: str | None = None) -> str:
    """
    Send an email to myself or to another recipient with a subject and content.

    Args:
        - subject: str - Text subject of the email
        - content: str - Text body content of the email
        - attachment_path: optional path to an image file saved by search_and_save_images
        - to_email: optional recipient email address. If omitted, uses the default in send_mail.
    """

    try:
        # Only pass the to_email argument to send_mail when provided so the send_mail
        # default recipient remains in effect when not specified.
        if to_email:
            send_mail(to_email=to_email, subject=subject, content=content,
                      attachment_path=attachment_path)
        else:
            send_mail(subject=subject, content=content,
                      attachment_path=attachment_path)
    except Exception as e:
        return f"Not Sent: {e}"
    return "Sent email"


@tool
def get_unread_emails(hours: int = 24) -> str:
    """
    Retreive emails that are unread within last N hours

    Args:
        - hours: int = 24 - number of hours ago to retreive from inbox

    Returns:
        - a string of emails separated by line "----"
    """

    try:
        emails = read_inbox(hours_ago=hours)
    except:
        return "Error getting unread emails"

    cleaned = []
    for email in emails:
        print(email)
        data = email.copy()
        if "html_body" in data:
            data.pop('html_body')
        msg = ""
        for k, v in data.items():
            msg += f"{k}:\t{v}\n"
        cleaned.append(msg)

    return "\n----\n".join(cleaned)


'''
This cleaning and returning everything as string is essential as LLMs are very good at
reading and reasoning over texts compared to dictionaries. Also this can be thought of
as when using an API you might need to transform data into a particular format before
passing it to the LLM.
'''


@tool
def search_and_save_images(topic: str, num_images: int = 3) -> str:
    """
    Search Wikimedia Commons for images related to a topic and save them
    to the images folder.

    Args:
        topic: Topic to search for.
        num_images: Number of images to download.

    Returns:
        Paths of the downloaded images.
    """

    try:
        search_url = "https://commons.wikimedia.org/w/api.php"

        params = {
            "action": "query",
            "generator": "search",
            "gsrsearch": topic,
            "gsrnamespace": 6,
            "gsrlimit": num_images,
            "prop": "imageinfo",
            "iiprop": "url",
            "iiurlwidth": 1200,
            "format": "json",
        }

        response = requests.get(
            search_url,
            params=params,
            headers={"User-Agent": "DockerLearn/1.0"},
            timeout=20,
        )
        response.raise_for_status()

        pages = response.json().get("query", {}).get("pages", {})

        if not pages:
            return f"No images found for: {topic}"

        image_dir = Path("/app/images")
        image_dir.mkdir(parents=True, exist_ok=True)

        saved_paths = []

        for i, page in enumerate(pages.values(), start=1):
            image_info = page.get("imageinfo", [])

            if not image_info:
                continue

            image_url = (
                image_info[0].get("thumburl")
                or image_info[0].get("url")
            )

            if not image_url:
                continue

            image_path = image_dir / f"{topic.replace(' ', '_')}_{i}.jpg"

            image_response = requests.get(
                image_url,
                headers={"User-Agent": "DockerLearn/1.0"},
                timeout=30,
            )
            image_response.raise_for_status()

            image_path.write_bytes(image_response.content)

            saved_paths.append(str(image_path))

        if not saved_paths:
            return f"Could not download images for: {topic}"

        return "\n".join(saved_paths)

    except Exception as e:
        return f"Failed to download images: {e}"


@tool
def web_search(query: str) -> list[Source]:
    """
    Web search tool for searching the web about a topic

    Args:
        - query: the query around topic to search with

    Returns: 
        - A list of search items
    """
    print(
        f"[TOOL CALL] web_search: {query}")  # cheap alternative to logging internal tool calls lol
    client = Parallel(api_key=PARALLEL_API_KEY)
    query_rewriter_llm = get_openai_llm().with_structured_output(SearchQueries)

    messages = [
        ("system",
         """You are a query rewriter and you take queries and make them comprehensive
        Understand the query and return at least 3 comprehensive rewritten queries."""),
        ("human", f"{query}, do not use markdown in your response only plain text")]

    result = query_rewriter_llm.invoke(messages)

    response = client.search(
        objective="Find latest information about search queries.",
        search_queries=result.queries,
        mode="advanced",
        advanced_settings={"max_results": 10},
    ).model_dump()

    sources: list[Source] = []
    for result in response.get("results", []):
        url = result.get("url", "")
        title = result.get("title", "")
        excerpts = result.get("excerpts") or []
        excerpt_text = " ".join(excerpts)[:500]

        sources.append(
            Source(
                title=title,
                url=url,
                publisher=_publisher_from_url(url),
                excerpt=excerpt_text
            )
        )

    if current_sources.get() is not None:
        current_sources.get().extend(sources)

    return sources


@tool
def search_arxiv(query: str, max_results: int = 5) -> list[Source]:
    """
    Search arXiv and return relevant paper details.

    Args: 
        - query: relevant query to research topic used to search arxiv 
        - max_results: maximum number of papers to retrieve
    Returns:
        - a list of Source objects
    """
    print(
        f"[TOOL CALL] search_arxiv: {query}")  # cheaep alternative to logging internal tool calls
    client = arxiv.Client()

    search = arxiv.Search(
        query=query,
        max_results=max_results,
        sort_by=arxiv.SortCriterion.Relevance,
    )

    sources: list[Source] = []

    sources: list[Source] = []

    for r in client.results(search):
        summary = " ".join(r.summary.split())

        sources.append(
            Source(
                title=r.title,
                url=r.entry_id,
                publisher="arXiv",
                excerpt=summary[:1000],
            )
        )

    print(f"[TOOL RESULT] search_arxiv: {len(sources)} results")

    if current_sources.get() is not None:
        current_sources.get().extend(sources)

    return sources


@tool
def generate_report_docx(result: ResearchResult | dict, title: str = "Research Report") -> str:
    """
    Generate a Word (.docx) report from a ResearchResult and save it to disk.

    Args:
        - result: ResearchResult or dict with keys (answer, sources, confidence)
        - title: title to put at the top of the document

    Returns:
        - Path to the saved .docx file
    """
    if Document is None:
        return (
            "python-docx is required to generate reports. Install it with: "
            "pip install python-docx"
        )

    # Normalize to ResearchResult
    if isinstance(result, dict):
        res = ResearchResult(**result)
    elif isinstance(result, ResearchResult):
        res = result
    else:
        # try to coerce
        try:
            res = ResearchResult.model_validate(result)  # pydantic v2 shape
        except Exception:
            try:
                res = ResearchResult(**result)
            except Exception as e:
                return f"Invalid ResearchResult provided: {e}"

    reports_dir = Path("/app/reports")
    reports_dir.mkdir(parents=True, exist_ok=True)

    safe_title = title.replace("/", "_").replace("\\\\", "_")
    filename = f"{safe_title[:50].replace(' ', '_')}.docx"
    filepath = reports_dir / filename

    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(res.answer)

    doc.add_heading("Sources", level=2)
    for s in res.sources:
        # Each source may be a Source model or a dict
        if hasattr(s, "title"):
            stitle = s.title
            surl = s.url
            spub = getattr(s, "publisher", "")
            srel = getattr(s, "relevance", "")
        else:
            stitle = s.get("title", "")
            surl = s.get("url", "")
            spub = s.get("publisher", "")
            srel = s.get("relevance", "")

        p = doc.add_paragraph()
        p.add_run(stitle).bold = True
        if surl:
            p.add_run(f" — {surl}")
        if spub:
            p.add_run(f"\nPublisher: {spub}")
        if srel:
            p.add_run(f"\nRelevance: {srel}")

        doc.add_paragraph()

    doc.add_heading("Confidence", level=2)
    doc.add_paragraph(res.confidence)

    try:
        doc.save(filepath)
    except Exception as e:
        return f"Failed to save report: {e}"

    return str(filepath)
